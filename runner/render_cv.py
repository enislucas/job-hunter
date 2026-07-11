"""Job Hunter render chain.

Usage: py runner\\render_cv.py <package-folder>
Renders every top-level .md whose name is cv.md or letter.md in the folder to
<name>.docx (python-docx) and <name>.pdf (single-column HTML printed by Edge headless).

Markdown subset understood (keep sources inside it):
  # Name            -> document title line (bold, larger)
  plain lines directly after the title, before the first ## -> contact block
  ## Section        -> section heading (bold, small caps feel via size)
  - bullet          -> bullet paragraph
  **bold** inline   -> bold run
  everything else   -> normal paragraph
No tables, no images, one column, by design (ATS rules A1).
"""
import sys, re, subprocess, html
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

EDGE_PATHS = [
    r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
]

BODY_FONT = "Calibri"


def parse(md_text):
    """Return list of (kind, text) tuples: title, contact, heading, bullet, para."""
    out = []
    lines = md_text.splitlines()
    seen_title = False
    seen_heading = False
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# ") and not seen_title:
            out.append(("title", line[2:].strip()))
            seen_title = True
        elif line.startswith("## "):
            out.append(("heading", line[3:].strip()))
            seen_heading = True
        elif line.startswith("- "):
            out.append(("bullet", line[2:].strip()))
        else:
            kind = "contact" if (seen_title and not seen_heading) else "para"
            out.append((kind, line.strip()))
    return out


def add_runs(paragraph, text, size, bold_all=False):
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            if bold_all:
                run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(size)


def to_docx(items, path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)
    for kind, text in items:
        if kind == "title":
            p = doc.add_paragraph()
            add_runs(p, text, 17, bold_all=True)
            p.paragraph_format.space_after = Pt(2)
        elif kind == "contact":
            p = doc.add_paragraph()
            add_runs(p, text, 9.5)
            p.paragraph_format.space_after = Pt(1)
        elif kind == "heading":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text.upper())
            run.bold = True
            run.font.name = BODY_FONT
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, text, 10.5)
            p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph()
            add_runs(p, text, 10.5)
            p.paragraph_format.space_after = Pt(3)
    doc.save(path)


def md_inline_html(text):
    text = html.escape(text)
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)


def to_html(items):
    body = []
    for kind, text in items:
        t = md_inline_html(text)
        if kind == "title":
            body.append(f"<h1>{t}</h1>")
        elif kind == "contact":
            body.append(f"<p class='c'>{t}</p>")
        elif kind == "heading":
            body.append(f"<h2>{t}</h2>")
        elif kind == "bullet":
            body.append(f"<li>{t}</li>")
        else:
            body.append(f"<p>{t}</p>")
    # wrap consecutive <li> into <ul>
    joined = []
    in_ul = False
    for b in body:
        if b.startswith("<li>") and not in_ul:
            joined.append("<ul>")
            in_ul = True
        if not b.startswith("<li>") and in_ul:
            joined.append("</ul>")
            in_ul = False
        joined.append(b)
    if in_ul:
        joined.append("</ul>")
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'><style>"
        "@page{size:A4;margin:1.4cm 1.7cm}"
        "body{font-family:Calibri,'Segoe UI',sans-serif;font-size:10.5pt;color:#111;line-height:1.38;margin:0}"
        "h1{font-size:17pt;margin:0 0 2pt}"
        "p.c{font-size:9.5pt;margin:0 0 1pt;color:#222}"
        "h2{font-size:10.5pt;text-transform:uppercase;letter-spacing:.06em;color:#1f3a5f;"
        "border-bottom:1pt solid #c9d4e2;padding-bottom:2pt;margin:11pt 0 4pt}"
        "p{margin:0 0 3.5pt}ul{margin:0 0 4pt 14pt;padding:0}li{margin:0 0 2.5pt}"
        "</style></head><body>" + "".join(joined) + "</body></html>"
    )


def edge():
    for p in EDGE_PATHS:
        if Path(p).exists():
            return p
    raise SystemExit("msedge.exe not found in known locations")


def render(folder):
    folder = Path(folder)
    done = []
    for name in ("cv", "letter"):
        src = folder / f"{name}.md"
        if not src.exists():
            continue
        items = parse(src.read_text(encoding="utf-8-sig"))
        to_docx(items, folder / f"{name}.docx")
        html_path = folder / f"_{name}_print.html"
        html_path.write_text(to_html(items), encoding="utf-8")
        pdf_path = folder / f"{name}.pdf"
        subprocess.run(
            [edge(), "--headless", "--disable-gpu", "--no-first-run",
             f"--print-to-pdf={pdf_path}", "--no-pdf-header-footer",
             html_path.resolve().as_uri()],
            check=True, timeout=60,
        )
        html_path.unlink()
        done.append(name)
    print("rendered:", ", ".join(done) if done else "nothing (no cv.md/letter.md found)")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit(__doc__)
    render(sys.argv[1])
